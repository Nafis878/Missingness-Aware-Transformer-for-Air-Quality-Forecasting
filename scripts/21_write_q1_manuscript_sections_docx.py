from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs"
TABLE_DIR = OUT_DIR / "tables"
DOCX_PATH = OUT_DIR / "Q1_journal_manuscript_sections.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_FILL = "F4F6F9"
BORDER = "B8C7D9"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def set_run_font(run, size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = "Calibri"
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc: Document, text: str = "", style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=9, italic=True, color="555555")
    return p


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.333
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    footer_run = footer.add_run("Manuscript draft sections")
    set_run_font(footer_run, size=9, color="666666")


def format_table(table, widths: list[int]) -> None:
    set_table_width(table, widths)
    set_table_borders(table)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.167
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, size=8.8)
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, size=8.8, bold=True, color=INK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    format_table(table, widths)
    return table


def add_lead_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = text
    set_cell_shading(cell, LIGHT_FILL)
    format_table(table, [9360])
    for p in cell.paragraphs:
        for run in p.runs:
            set_run_font(run, size=10.5, bold=True, color=INK)
    add_para(doc, "")


def pct(value: str) -> str:
    return f"{float(value):.2f}"


def rmse(value: str) -> str:
    return f"{float(value):.2f}"


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    final_rows = read_csv(TABLE_DIR / "adaptive_sota_portfolio_trainval_vs_table_best.csv")
    paired_rows = read_csv(TABLE_DIR / "adaptive_sota_portfolio_trainval_paired_tests.csv")
    core_rows = read_csv(TABLE_DIR / "core_model_before_after_imputation_forecasting.csv")
    availability_rows = read_csv(TABLE_DIR / "core_model_before_after_imputation_availability.csv")

    improvements = [float(row["relative_improvement_pct"]) for row in final_rows]
    mean_improvement = sum(improvements) / len(improvements)
    min_improvement = min(improvements)
    max_improvement = max(improvements)
    strict_wins = sum(row["significant_win"] == "True" for row in paired_rows)
    directional_wins = sum(row["directional_win"] == "True" for row in paired_rows)

    dhaka_core = [row for row in core_rows if row["dataset"] == "Dhaka"]
    sarima_improved = [row for row in dhaka_core if row["model"] == "SARIMA" and float(row["improvement_pct"]) > 0]
    dlinear_worse = [row for row in dhaka_core if row["model"] == "DLinear" and float(row["improvement_pct"]) < 0]
    gru_worse = [row for row in dhaka_core if row["model"] == "GRU" and float(row["improvement_pct"]) < 0]
    missing_pairs = [row for row in availability_rows if row["before_after_pair_available"] == "False"]

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    tr = title.add_run("Manuscript Draft Sections")
    set_run_font(tr, size=20, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    sr = subtitle.add_run(
        "Missingness-aware imputation and adaptive forecasting for PM2.5 air-quality prediction"
    )
    set_run_font(sr, size=12, italic=True, color="555555")

    add_lead_callout(
        doc,
        "Honest claim used in this draft: the adaptive portfolio improves RMSE in all 9 "
        "dataset-horizon comparisons, and the global same-direction pattern is statistically "
        "supported. The current evidence does not support claiming independent statistical "
        "significance for every individual cell."
    )

    doc.add_heading("Abstract", level=1)
    add_para(
        doc,
        "Air-quality forecasting is complicated by irregular and often non-random missingness, "
        "especially in PM2.5 monitoring data from dense urban environments. This study evaluates "
        "whether a merged hybrid_top8 imputation strategy, combined with missingness-aware model "
        "selection, improves multi-horizon PM2.5 forecasting across three city-level datasets: "
        "Dhaka, Delhi, and Beijing. The forecasting task was evaluated at 6, 24, and 72 hour horizons "
        "using chronological train-validation-test splits and paired test-set comparisons. The model "
        "pool included statistical baselines, ARIMA/SARIMA, recurrent neural networks, DLinear, "
        "Transformer/MAT variants, tabular tree ensembles, and validation-selected ensembles. "
        "Rather than claiming that one imputation method improves every standalone backbone, the "
        "proposed final system uses validation evidence to select or blend the strongest model "
        "configuration for each dataset-horizon setting.",
    )
    add_para(
        doc,
        f"The resulting adaptive portfolio reduced RMSE relative to the previous best result in "
        f"all 9 dataset-horizon comparisons, with relative gains ranging from {min_improvement:.2f}% "
        f"to {max_improvement:.2f}% and a mean gain of {mean_improvement:.2f}%. The universal "
        "directional pattern was statistically supported by an exact one-sided sign test "
        "(p = 0.001953), a one-sided Wilcoxon signed-rank test (p = 0.001953), Fisher-combined "
        "one-sided Diebold-Mariano evidence (p = 1.52841e-05), and weighted Stouffer-combined "
        "evidence (p = 3.85672e-05). However, under a conservative cell-level criterion requiring "
        f"a negative RMSE difference, a paired-bootstrap confidence interval below zero, and "
        f"DM p < 0.05, only {strict_wins} of 9 individual cells were independently significant. "
        "The findings therefore support a defensible global superiority claim for the adaptive "
        "missingness-aware forecasting portfolio on the tested benchmarks, while showing that "
        "direct imputation benefits are model- and horizon-dependent."
    )

    doc.add_heading("Introduction", level=1)
    add_para(
        doc,
        "Fine-grained PM2.5 forecasting is a core requirement for air-quality early warning, "
        "public-health planning, and pollution exposure assessment. In practice, forecasting "
        "systems must operate on monitoring streams that contain sensor outages, maintenance gaps, "
        "communication failures, calibration anomalies, and weather-dependent observation patterns. "
        "Missing values are therefore not a peripheral preprocessing inconvenience; they are part "
        "of the data-generating environment that the forecaster must learn to handle."
    )
    add_para(
        doc,
        "A common response is to impute missing observations before training a forecasting model. "
        "Simple imputers such as forward fill, interpolation, mean replacement, or KNN imputation "
        "can stabilize the input matrix, but they may also remove information carried by the "
        "missingness pattern itself. Deep forecasting models face a related problem: once values "
        "are filled, the model may treat observed and reconstructed values as equally reliable "
        "unless explicit masks or missingness-aware features are supplied. This is especially "
        "important for air-quality data, where long gaps and seasonally structured missingness can "
        "alter the apparent dynamics of PM2.5 concentration."
    )
    add_para(
        doc,
        "The present work addresses this issue by treating imputation and forecasting as a linked "
        "model-selection problem rather than as a fixed preprocessing step. The study first "
        "constructs a merged hybrid_top8 imputation representation from the strongest imputation "
        "candidates available in the experimental benchmark. It then evaluates a broad model pool "
        "that includes statistical models, recurrent networks, linear sequence models, and "
        "Transformer-style missingness-aware architectures. Finally, an adaptive portfolio selects "
        "or blends models using validation performance, and the selected configuration is assessed "
        "on held-out test periods."
    )
    add_para(
        doc,
        "The central contribution is not the assertion that hybrid imputation universally improves "
        "every individual model. The experiments show that such a statement would be inaccurate: "
        "some backbones improve, while others degrade after direct imputation. The contribution is "
        "instead an empirically validated adaptive framework that uses imputation, masks, model "
        "diversity, and validation-based selection to obtain consistent test-set gains across all "
        "evaluated city-horizon combinations. This framing is important for a Q1 journal submission "
        "because it distinguishes a robust global result from an overgeneralized preprocessing claim."
    )

    doc.add_heading("Methodology", level=1)
    doc.add_heading("Data and Forecasting Task", level=2)
    add_para(
        doc,
        "The experiments used PM2.5 time-series datasets for Dhaka, Delhi, and Beijing. Each dataset "
        "was evaluated under a chronological forecasting protocol to avoid temporal leakage. Models "
        "were trained on historical observations, tuned or selected using a validation period, and "
        "finally evaluated on a held-out test period. Forecasting performance was measured at 6, 24, "
        "and 72 hour horizons, which represent short-, medium-, and longer-range operational use cases."
    )
    add_para(
        doc,
        "The primary metric reported in the comparative results is RMSE, because it penalizes large "
        "forecasting errors and is widely used in PM2.5 forecasting benchmarks. MAE and R2 were also "
        "tracked where available, but the final cross-model claims are based on RMSE to maintain a "
        "single, consistent ranking criterion."
    )

    doc.add_heading("Merged hybrid_top8 Imputation", level=2)
    add_para(
        doc,
        "The imputation stage used the merged hybrid_top8 representation derived from the strongest "
        "imputation candidates in the benchmark outputs. The objective was to recover a complete "
        "forecasting input while retaining the information that values were originally missing. "
        "For models able to use missingness indicators, imputed values were paired with masks or "
        "missingness-aware features so that reconstructed observations were not treated identically "
        "to directly observed measurements."
    )
    add_para(
        doc,
        "The imputed representation was evaluated in two ways. First, direct before-after comparisons "
        "were performed for all model families with available paired artifacts. Second, hybrid_top8 "
        "was included in the broader adaptive forecasting portfolio, where its value was judged by "
        "validation-selected downstream forecast accuracy rather than by imputation error alone."
    )

    doc.add_heading("Forecasting Model Pool", level=2)
    add_para(
        doc,
        "The implemented model pool covered both classical and neural forecasting families: "
        "persistence and seasonal baselines, ARIMA/SARIMA, LSTM, GRU, GRU-D, DLinear, PatchTST when "
        "available, Transformer/MAT variants, missingness-dropout variants, tabular ExtraTrees "
        "models, and ensemble configurations. This breadth was included to test whether missingness "
        "handling transfers across modeling assumptions rather than only within a single neural "
        "architecture."
    )
    add_para(
        doc,
        "For the final reported system, validation performance determined the selected portfolio "
        "member for each dataset-horizon pair. Several final winners were tabular ExtraTrees models "
        "refit on train plus validation data, while other cells were best served by validation-selected "
        "ensembles or a linear-tabular blend. This adaptive design reflects the empirical observation "
        "that no single standalone backbone dominated all datasets and horizons."
    )

    doc.add_heading("Evaluation and Statistical Testing", level=2)
    add_para(
        doc,
        "The final comparison was made against the previous best RMSE available for each dataset and "
        "horizon. Paired test-set errors were used where prediction bundles were available. For each "
        "cell, the analysis computed the RMSE difference between the adaptive portfolio and the "
        "previous best model, paired-bootstrap confidence intervals, and Diebold-Mariano statistics. "
        "A strict individual-cell win required three conditions: the RMSE difference had to be below "
        "zero, the 95% paired-bootstrap confidence interval had to remain below zero, and the "
        "Diebold-Mariano p-value had to be below 0.05."
    )
    add_para(
        doc,
        "Because the paper's strongest claim concerns the repeated direction of improvement across "
        "datasets and horizons, the study also tested the global pattern. The exact one-sided sign "
        "test and one-sided Wilcoxon signed-rank test were applied to the nine RMSE deltas. In "
        "addition, one-sided Diebold-Mariano p-values were combined using Fisher's method and a "
        "weighted Stouffer method. This two-layer statistical design separates cell-level evidence "
        "from global evidence and prevents overstating the certainty of small-margin individual cells."
    )

    doc.add_heading("Results", level=1)
    doc.add_heading("Overall Adaptive Portfolio Performance", level=2)
    add_para(
        doc,
        f"The adaptive portfolio improved upon the previous best RMSE in every evaluated "
        f"dataset-horizon combination ({directional_wins}/9). Relative RMSE improvement ranged from "
        f"{min_improvement:.2f}% for Beijing at the 72 hour horizon to {max_improvement:.2f}% for "
        f"Delhi at the 24 hour horizon, with an average improvement of {mean_improvement:.2f}% "
        "across the nine comparisons. The results are summarized in Table 1."
    )

    result_rows = [
        [
            row["dataset"],
            row["horizon"],
            row["adaptive_portfolio_method"].replace("_", " "),
            rmse(row["previous_table_best_RMSE"]),
            rmse(row["RMSE"]),
            pct(row["relative_improvement_pct"]),
        ]
        for row in final_rows
    ]
    add_caption(doc, "Table 1. Adaptive portfolio RMSE against the previous best result.")
    add_table(
        doc,
        ["Dataset", "Horizon", "Selected portfolio member", "Previous best RMSE", "Portfolio RMSE", "Improvement (%)"],
        result_rows,
        [920, 820, 3140, 1450, 1450, 1580],
    )

    doc.add_heading("Statistical Evidence", level=2)
    add_para(
        doc,
        "The cell-level and global statistical tests lead to a nuanced but strong conclusion. At the "
        "cell level, all nine RMSE differences were in the favorable direction, but only two met the "
        "strict individual-cell significance rule. At the global level, the repeated 9/9 improvement "
        "pattern was unlikely under a no-improvement null, and the combined paired evidence also "
        "rejected the global null. These results support a universal directional superiority claim "
        "for the adaptive portfolio on the tested benchmark suite, but they do not support the claim "
        "that every individual cell is independently significant."
    )
    add_caption(doc, "Table 2. Statistical support for the final adaptive portfolio claim.")
    add_table(
        doc,
        ["Evidence", "Result", "Publication interpretation"],
        [
            ["Directional wins", "9/9", "Every dataset-horizon cell improved in RMSE."],
            ["Strict cell-level wins", f"{strict_wins}/9", "Only two cells satisfy the conservative individual-cell rule."],
            ["Exact one-sided sign test", "p = 0.001953", "The 9/9 same-direction pattern is globally significant."],
            ["One-sided Wilcoxon signed-rank", "p = 0.001953", "Magnitude-ranked paired deltas support a global improvement."],
            ["Fisher combined one-sided DM", "p = 1.52841e-05", "Combined paired forecast-error evidence rejects the global null."],
            ["Weighted Stouffer one-sided DM", "p = 3.85672e-05", "Weighted combined evidence also supports global superiority."],
        ],
        [2300, 2200, 4860],
    )

    doc.add_heading("Effect of Imputation Across Implemented Models", level=2)
    add_para(
        doc,
        "The direct before-after imputation analysis shows that hybrid_top8 imputation is not a "
        "uniformly beneficial preprocessing step for every implemented model. This is an important "
        "negative result and should be retained in the manuscript. In Dhaka, SARIMA improved across "
        f"all three horizons after hybrid_top8 imputation ({len(sarima_improved)}/3), whereas DLinear "
        f"and standard GRU degraded across all three available horizons ({len(dlinear_worse)}/3 and "
        f"{len(gru_worse)}/3, respectively). GRU-D improved at the 6 hour horizon but degraded at "
        "24 and 72 hours. The Transformer/MAT family showed mixed effects across datasets: for "
        "example, Delhi MAT improved at 72 hours but worsened at 6 and 24 hours, while Beijing "
        "missingness-dropout MAT improved at 6 hours but worsened at longer horizons."
    )
    add_para(
        doc,
        f"Before-after paired artifacts were available for Dhaka across ARIMA/SARIMA, LSTM, GRU, "
        "GRU-D, DLinear, and MAT-family models, and for Delhi and Beijing within the MAT-family "
        f"experiments. They were not available for {len(missing_pairs)} listed dataset-model "
        "combinations, including several non-MAT families in Delhi and Beijing and PatchTST in the "
        "current saved outputs. The manuscript should therefore avoid implying that every implemented "
        "family has complete before-after coverage across all three cities."
    )
    add_caption(doc, "Table 3. Honest interpretation of direct hybrid_top8 imputation effects.")
    add_table(
        doc,
        ["Finding", "Evidence from saved artifacts", "Interpretation"],
        [
            [
                "SARIMA benefits on Dhaka",
                "RMSE improves at 6, 24, and 72 hour horizons.",
                "Classical seasonal modeling can benefit from the merged imputed representation in this dataset.",
            ],
            [
                "Some neural and linear backbones degrade",
                "Dhaka DLinear and GRU worsen at all three available horizons.",
                "Imputation alone is not sufficient for a universal model-level claim.",
            ],
            [
                "MAT-family behavior is mixed",
                "Delhi improves at 72 hours but not at 6 or 24 hours; Beijing gains are horizon-dependent.",
                "Masks and missingness-aware features help in some regimes but require validation selection.",
            ],
            [
                "Adaptive portfolio is the defensible final claim",
                "Final selected system beats the previous best in all 9 city-horizon cells.",
                "The Q1-ready claim should be portfolio-level, not blanket imputation-level superiority.",
            ],
        ],
        [1900, 3600, 3860],
    )

    doc.add_heading("Recommended Manuscript Claim", level=2)
    add_para(
        doc,
        "A defensible wording for the Results section is: The adaptive missingness-aware forecasting "
        "portfolio achieved lower RMSE than the previous best result in all nine dataset-horizon "
        "comparisons across Dhaka, Delhi, and Beijing. The 9/9 same-direction pattern was significant "
        "under exact sign and Wilcoxon signed-rank testing, and combined paired Diebold-Mariano "
        "evidence also rejected the global no-improvement null. However, only two individual cells "
        "met the strict cell-level significance criterion, so the evidence supports global "
        "benchmark-suite superiority rather than independent significance in every individual "
        "dataset-horizon cell."
    )
    add_para(
        doc,
        "For publication, the figures that best support this narrative are the universal RMSE "
        "comparison, relative improvement heatmap, paired delta forest plot, universal statistical "
        "support panel, and the all-model before-after imputation figures already saved in the "
        "outputs/figures directory. Together, these figures show both the strength of the final "
        "portfolio result and the necessary honesty about mixed direct imputation effects."
    )

    # Add a final blank paragraph to avoid cramped footer interaction on the last page.
    add_para(doc, "")
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
