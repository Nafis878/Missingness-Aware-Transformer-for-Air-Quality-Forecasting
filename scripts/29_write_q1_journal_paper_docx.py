"""Write the Q1 journal manuscript DOCX for the final MAT ensemble study."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs"
TABLE_DIR = OUT_DIR / "tables"
FIG_DIR = OUT_DIR / "figures"
DOCX_PATH = OUT_DIR / "Q1_MAT_Ensemble_Journal_Paper.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
BORDER = "B8C7D9"
MUTED = "555555"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def set_run_font(run, size=None, bold=None, italic=None, color=None) -> None:
    run.font.name = "Calibri"
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.333
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Q1 manuscript draft")
    set_run_font(run, size=9, color="666666")


def para(doc: Document, text: str = "", style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run)
    return p


def caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=9, italic=True, color=MUTED)
    return p


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    t = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_width(t, widths)
    set_table_borders(t)
    for ridx, row in enumerate(t.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(cell, HEADER_FILL if ridx == 0 else "FFFFFF")
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.167
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, size=8.7, bold=(ridx == 0), color=(INK if ridx == 0 else None))
    return t


def callout(doc: Document, text: str, label: str = "Key claim") -> None:
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    cell.text = f"{label}: {text}"
    set_cell_shading(cell, LIGHT_FILL)
    set_table_width(t, [9360])
    set_table_borders(t)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            set_run_font(run, size=10.2, bold=True, color=INK)
    para(doc, "")


def add_figure(doc: Document, image: Path, caption_text: str, width: float = 6.25) -> None:
    if not image.exists():
        para(doc, f"[Missing figure: {image.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(image), width=Inches(width))
    caption(doc, caption_text)


def fmt(value: str | float, digits=2) -> str:
    return f"{float(value):.{digits}f}"


def get_row(rows: list[dict[str, str]], model: str, horizon: int) -> dict[str, str]:
    for row in rows:
        if row["model"] == model and int(row["horizon"]) == horizon:
            return row
    raise KeyError((model, horizon))


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    comp = read_csv(TABLE_DIR / "final_mat_ensemble_comparison_summary.csv")
    ablation = read_csv(TABLE_DIR / "ablation_metrics_after_imputation.csv")
    sig = read_csv(TABLE_DIR / "combined_seed_significance_validation_convex_intercept_stack.csv")

    doc = Document()
    style_document(doc)

    title = (
        "Validation-Calibrated Missingness-Aware Transformer Ensemble for Robust "
        "Multi-Horizon PM2.5 Forecasting under Incomplete Air-Quality Monitoring"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    set_run_font(r, size=19, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Manuscript draft prepared for Q1 journal submission")
    set_run_font(r, size=11, italic=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Author: [Your Name]  |  Affiliation: [Your Institution]")
    set_run_font(r, size=10.5, color=MUTED)

    callout(
        doc,
        "The final validation-calibrated MAT ensemble achieves the lowest RMSE at H6, H24, "
        "and H72, improves over Vanilla Transformer by 2.83, 4.08, and 4.28 RMSE points, "
        "and passes 42/42 combined seed-level Diebold-Mariano significance checks after "
        "Holm correction.",
        "Main result",
    )

    doc.add_heading("Abstract", level=1)
    para(
        doc,
        "Accurate PM2.5 forecasting is essential for air-quality early warning, exposure "
        "assessment, and public-health decision support. However, real-world air-quality "
        "monitoring networks are frequently incomplete because of sensor outages, maintenance "
        "interruptions, communication failures, and weather-dependent measurement gaps. This "
        "paper proposes and evaluates a validation-calibrated Missingness-Aware Transformer "
        "(MAT) ensemble for robust multi-horizon PM2.5 forecasting under incomplete monitoring "
        "conditions. The framework combines imputed and missingness-aware prediction pathways, "
        "uses validation-calibrated convex-intercept stacking, and evaluates forecast accuracy "
        "at 6, 24, and 72 hour horizons. The final model is compared with recurrent networks, "
        "GRU-D, DLinear, PatchTST, two-stage impute-then-Transformer pipelines using KNN, MICE, "
        "and SAITS, Vanilla Transformer, Full MAT, and multiple MAT variants."
    )
    para(
        doc,
        "On the held-out test set, the final validation-calibrated MAT ensemble obtains RMSE "
        "values of 65.78, 74.23, and 77.55 at H6, H24, and H72, respectively. It reduces RMSE "
        "relative to the Vanilla Transformer by 2.83, 4.08, and 4.28 points and improves over "
        "the Variant B dual-input ridge model by 1.50, 0.84, and 1.78 points. Combined seed-level "
        "Diebold-Mariano testing with Holm correction shows statistically significant superiority "
        "in all 42 model-horizon comparisons. These findings indicate that explicitly combining "
        "missingness-aware representation with validation-calibrated ensemble correction yields "
        "a more reliable forecasting system than either imputation-only or standalone Transformer "
        "modeling."
    )
    p = para(doc)
    r = p.add_run("Keywords: ")
    set_run_font(r, bold=True)
    r = p.add_run(
        "PM2.5 forecasting; missing data; Missingness-Aware Transformer; imputation; "
        "time-series forecasting; validation-calibrated ensemble; Diebold-Mariano test"
    )
    set_run_font(r)

    doc.add_heading("1. Introduction", level=1)
    para(
        doc,
        "Air pollution remains a major environmental and public-health challenge, particularly "
        "in rapidly urbanizing regions where particulate matter concentrations fluctuate with "
        "traffic, industry, meteorology, regional transport, and seasonal emissions. Forecasting "
        "PM2.5 several hours to several days ahead can support early warnings, mitigation actions, "
        "and exposure reduction. Yet operational forecasting systems must learn from monitoring "
        "streams that are rarely complete. Missingness arises from sensor faults, instrument "
        "maintenance, calibration issues, power interruptions, data communication gaps, and quality "
        "control removal of physically implausible values."
    )
    para(
        doc,
        "The common two-stage response is to first impute missing observations and then train a "
        "forecasting model on the completed matrix. This procedure stabilizes model input but can "
        "also obscure the reliability of reconstructed values. A model receiving only imputed values "
        "may fail to distinguish a directly observed concentration from a value estimated from a "
        "long outage. Missingness-aware models address this limitation by supplying the observation "
        "mask, time-gap information, or learned missingness embeddings to the forecaster."
    )
    para(
        doc,
        "This paper develops that idea into a validation-calibrated MAT ensemble. The proposed "
        "system is motivated by a practical observation from the experiments: no single standalone "
        "backbone dominates across all horizons and all missingness regimes. Vanilla Transformers "
        "benefit from clean imputed inputs, while MAT variants exploit missingness indicators and "
        "masked attention. A calibrated ensemble can use validation evidence to combine these "
        "complementary signals without retraining the neural backbones."
    )
    para(
        doc,
        "The contributions are fourfold. First, the study evaluates PM2.5 forecasting under a "
        "strict chronological protocol with three forecast horizons. Second, it compares MAT "
        "against strong neural, linear, and impute-then-forecast baselines. Third, it introduces "
        "a validation-calibrated convex-intercept MAT ensemble that combines missingness-aware and "
        "imputation-driven predictions. Fourth, it provides paired statistical evidence using "
        "seed-level Diebold-Mariano tests combined across seeds with Holm correction."
    )

    doc.add_heading("2. Literature Review", level=1)
    doc.add_heading("2.1 Time-Series Forecasting and Transformers", level=2)
    para(
        doc,
        "The Transformer architecture introduced self-attention as an alternative to recurrent "
        "sequence modeling and has since become a central design for sequence learning [1]. For "
        "time-series forecasting, Transformers offer flexible long-range dependency modeling, but "
        "their use in continuous temporal data has been debated because self-attention does not "
        "inherently encode ordered temporal structure. This has motivated both Transformer-based "
        "forecasting variants such as PatchTST [5] and simpler linear alternatives such as DLinear "
        "and LTSF-Linear [4]."
    )
    para(
        doc,
        "In air-quality forecasting, recurrent networks such as LSTM and GRU remain common because "
        "pollutant concentrations exhibit strong autocorrelation, diurnal cycles, and meteorological "
        "dependence. However, recurrent models can struggle with long histories and irregularly "
        "missing inputs. Transformer-style encoders can process longer windows more directly, but "
        "they require careful treatment of missing values and temporal features."
    )
    doc.add_heading("2.2 Missing Data in Multivariate Time Series", level=2)
    para(
        doc,
        "Missingness in multivariate time series is often informative: the absence of a measurement "
        "can correlate with sensor condition, environmental events, or operational regimes. GRU-D "
        "is an early deep model that explicitly incorporates masks and time intervals for missing "
        "values [2]. BRITS treats imputed values as variables within a bidirectional recurrent graph "
        "and learns them through backpropagation [7]. SAITS uses self-attention blocks for time-series "
        "imputation and has become a strong deep imputation baseline [3]. Classical approaches such "
        "as KNN, MICE, interpolation, and seasonal models remain competitive in many practical cases, "
        "especially when missingness is structured or the series is strongly periodic."
    )
    para(
        doc,
        "The literature therefore suggests that neither imputation-only nor missingness-native "
        "forecasting is universally optimal. Imputation can reduce input sparsity but may smooth or "
        "bias the signal; missingness-aware modeling can preserve reliability information but may "
        "underperform when imputation accurately reconstructs the underlying process. The present "
        "work follows this evidence by combining both pathways through validation-calibrated stacking."
    )

    doc.add_heading("3. Methodology", level=1)
    doc.add_heading("3.1 Data, Windowing, and Forecast Horizons", level=2)
    para(
        doc,
        "The core dataset consists of hourly air-quality and meteorological observations from a "
        "severely incomplete monitoring network. The forecasting target is PM2.5, with additional "
        "pollutants and meteorological variables used as covariates where available. The data are "
        "processed using a chronological split: training data precede validation data, and the final "
        "test period is held out. This split avoids leakage from future observations into model "
        "selection. Each input window covers one week of historical data, and forecasts are evaluated "
        "at H6, H24, and H72."
    )
    para(
        doc,
        "All models are evaluated only on observed target values. Target masks are used so that "
        "forecast errors are not computed against missing ground truth. Learned models are trained "
        "with seeds 42, 43, and 44, and reported results aggregate across these seeds."
    )

    doc.add_heading("3.2 Missingness-Aware Transformer Variants", level=2)
    para(
        doc,
        "The MAT backbone extends a standard Transformer encoder by adding learned missingness "
        "embeddings to the value representation. The input representation combines value projection, "
        "missingness projection, temporal features, station embeddings, and positional encodings. "
        "Variant B further masks attention to steps where the primary target is unobserved, allowing "
        "the model to treat missing-target timesteps differently from observed-target timesteps."
    )
    para(
        doc,
        "After imputation, three related Transformer-family variants are important. The Vanilla "
        "Transformer uses imputed values without explicit missingness information. MAT Variant B "
        "uses missingness-aware attention. Variant B dual-input ridge combines the native Variant B "
        "pathway with a Vanilla-input Variant B pathway using validation-fitted ridge stacking. This "
        "dual-input design reduces the Vanilla Transformer gap but does not by itself become the "
        "overall winner against every model under the strict per-seed test."
    )

    doc.add_heading("3.3 Validation-Calibrated MAT Ensemble", level=2)
    para(
        doc,
        "The final model is a validation-calibrated convex-intercept ensemble. It is fitted only on "
        "validation predictions and then applied once to the test predictions. For each horizon and "
        "seed, the method learns nonnegative convex weights over a pool of candidate predictors and "
        "an intercept correction. Candidate members include MAT Variant B, Variant B with Vanilla "
        "input, Vanilla Transformer, missingness-dropout MAT, KNN/MICE/SAITS two-stage Transformers, "
        "DLinear, and GRU-D. The ensemble therefore remains anchored in MAT-family modeling while "
        "using validation evidence to correct residual weaknesses."
    )
    callout(
        doc,
        "The final paper should describe the winner as a validation-calibrated MAT ensemble. It "
        "should not claim that plain Variant B dual-input ridge alone beats every model with "
        "per-seed statistical significance.",
        "Important wording",
    )

    doc.add_heading("3.4 Statistical Testing", level=2)
    para(
        doc,
        "Model comparisons use paired prediction errors on aligned test windows. For each baseline, "
        "horizon, and seed, the Diebold-Mariano test compares squared forecast errors from the final "
        "ensemble and the baseline. RMSE differences are computed as RMSE(final) minus RMSE(baseline), "
        "so negative values favor the final model. Seed-level p-values are then combined using Fisher "
        "and Stouffer methods, and Holm correction is applied across all 42 model-horizon comparisons. "
        "A comparison is considered significant when the final model is directionally better in all "
        "seeds and the combined Holm-corrected p-value is below 0.05."
    )

    doc.add_heading("4. Results", level=1)
    doc.add_heading("4.1 Transformer-Family Ablation after Imputation", level=2)
    ab_models = [
        ("vanilla_transformer", "Vanilla Transformer"),
        ("mat_full", "Full MAT"),
        ("mat_variant_B", "MAT Variant B"),
        ("mat_variant_B_vanilla_input", "Variant B + Vanilla input"),
        ("mat_variant_B_dual_input_ridge", "Variant B dual-input ridge"),
        ("mat_miss_dropout", "MAT missingness dropout"),
    ]
    ab_rows = []
    for model, label in ab_models:
        vals = [get_row(ablation, model, h) for h in [6, 24, 72]]
        avg = sum(float(v["RMSE_mean"]) for v in vals) / 3.0
        ab_rows.append([label, fmt(vals[0]["RMSE_mean"]), fmt(vals[1]["RMSE_mean"]), fmt(vals[2]["RMSE_mean"]), fmt(avg)])
    caption(doc, "Table 1. After-imputation Transformer-family ablation results.")
    table(doc, ["Model", "H6 RMSE", "H24 RMSE", "H72 RMSE", "Average"], ab_rows, [3000, 1500, 1500, 1500, 1500])
    para(
        doc,
        "Variant B dual-input ridge is the strongest after-imputation Transformer-family model, "
        "with average RMSE 73.89 compared with 76.25 for the Vanilla Transformer. It significantly "
        "beats the Vanilla Transformer at H6, H24, and H72. However, broader all-model testing showed "
        "that it did not fully dominate every non-Transformer baseline, motivating the validation-"
        "calibrated ensemble."
    )

    doc.add_heading("4.2 Final Model Comparison", level=2)
    selected_models = [
        ("hybrid8_transformer", "Vanilla Transformer"),
        ("variant_B_dual_input_ridge", "Variant B dual-input ridge"),
        ("two_stage_knn", "KNN + Transformer"),
        ("two_stage_saits", "SAITS + Transformer"),
        ("proposed", "Full MAT"),
        ("validation_convex_intercept_stack", "Final MAT ensemble"),
    ]
    comp_rows = []
    for model, label in selected_models:
        vals = [get_row(comp, model, h) for h in [6, 24, 72]]
        comp_rows.append([label, fmt(vals[0]["RMSE"]), fmt(vals[1]["RMSE"]), fmt(vals[2]["RMSE"])])
    caption(doc, "Table 2. Final validation-calibrated MAT ensemble against key baselines.")
    table(doc, ["Model", "H6 RMSE", "H24 RMSE", "H72 RMSE"], comp_rows, [3840, 1840, 1840, 1840])
    para(
        doc,
        "The final ensemble obtains the best RMSE at all three horizons: 65.78 at H6, 74.23 at "
        "H24, and 77.55 at H72. Compared with the Vanilla Transformer, the reductions are 2.83, "
        "4.08, and 4.28 RMSE points. Compared with Variant B dual-input ridge, the reductions are "
        "1.50, 0.84, and 1.78 RMSE points."
    )
    add_figure(
        doc,
        FIG_DIR / "q1_final_mat_key_rmse_comparison.png",
        "Figure 1. RMSE comparison between the final validation-calibrated MAT ensemble and key baselines.",
    )

    doc.add_heading("4.3 Significance against All Models", level=2)
    hardest = [
        ("two_stage_knn", 6),
        ("two_stage_saits", 6),
        ("proposed_md", 6),
        ("hybrid8_transformer", 6),
        ("hybrid8_transformer", 24),
        ("hybrid8_transformer", 72),
    ]
    sig_rows = []
    for baseline, horizon in hardest:
        row = next(r for r in sig if r["baseline"] == baseline and int(r["horizon"]) == horizon)
        sig_rows.append([
            row["baseline"].replace("_", " "),
            f"H{horizon}",
            fmt(row["RMSE_diff_mean_candidate_minus_baseline"]),
            f"{float(row['combined_fisher_p_holm']):.3g}",
            "Yes" if row["combined_significant_fisher_holm"] == "True" else "No",
        ])
    caption(doc, "Table 3. Representative combined seed-level significance results.")
    table(doc, ["Baseline", "Horizon", "RMSE diff", "Holm p", "Significant"], sig_rows, [3000, 1100, 1600, 1600, 2060])
    para(
        doc,
        "The final ensemble is directionally better than every baseline in all seeds. Under combined "
        "seed-level Diebold-Mariano testing with Holm correction, all 42 model-horizon comparisons "
        "are significant. This resolves the earlier H6 weakness of the standalone Variant B dual-input "
        "ridge model without claiming unsupported per-seed significance for every individual seed."
    )
    add_figure(
        doc,
        FIG_DIR / "q1_final_mat_combined_significance_heatmap.png",
        "Figure 2. Combined seed-level Diebold-Mariano significance after Holm correction.",
        width=5.7,
    )
    add_figure(
        doc,
        FIG_DIR / "q1_final_mat_ensemble_summary_panel.png",
        "Figure 3. Summary panel showing main comparison, gains over Vanilla Transformer, parity scatter, and significance distribution.",
    )

    doc.add_heading("5. Discussion", level=1)
    para(
        doc,
        "The results show that the main limitation of standalone MAT Variant B is not a lack of "
        "missingness awareness but the instability of a single architecture across all horizons and "
        "seeds. Variant B dual-input ridge closes the Vanilla Transformer gap and becomes the strongest "
        "Transformer-family member after imputation, but H6 comparisons against KNN, SAITS, and "
        "missingness-dropout MAT remain difficult under the strictest per-seed criterion. The final "
        "validation-calibrated ensemble solves this issue by using validation evidence to combine the "
        "strengths of imputation-driven and mask-aware pathways."
    )
    para(
        doc,
        "This finding aligns with the broader missing-data literature: the optimal handling of "
        "missingness depends on both the missingness mechanism and the temporal reconstructability "
        "of the series. In highly structured settings, imputation can be strong; in noisy and "
        "irregular settings, missingness indicators preserve important reliability information. The "
        "ensemble result is therefore not merely a numerical improvement but a methodological point: "
        "forecasting under incomplete monitoring is best treated as a joint imputation, representation, "
        "and validation-calibration problem."
    )
    para(
        doc,
        "The study has limitations. First, the final model is an ensemble, so it is less interpretable "
        "than a single Transformer backbone. Second, the significance claim depends on combined seed-"
        "level evidence rather than requiring every individual seed p-value to be below 0.05. Third, "
        "the final calibration is fitted on validation predictions from available model outputs; "
        "future work should evaluate whether the same strategy transfers unchanged to additional "
        "cities, pollutants, and sensor networks."
    )

    doc.add_heading("6. Conclusion", level=1)
    para(
        doc,
        "This paper presented a validation-calibrated Missingness-Aware Transformer ensemble for "
        "multi-horizon PM2.5 forecasting under incomplete air-quality monitoring. The final ensemble "
        "achieved RMSE values of 65.78, 74.23, and 77.55 at H6, H24, and H72, outperforming the "
        "Vanilla Transformer, standalone MAT variants, two-stage imputation Transformers, recurrent "
        "networks, DLinear, and PatchTST. Combined seed-level Diebold-Mariano testing with Holm "
        "correction confirmed statistically significant superiority across all 42 model-horizon "
        "comparisons."
    )
    para(
        doc,
        "The central conclusion is that missingness-aware forecasting benefits from validation-"
        "calibrated integration rather than relying on a single architecture or a single imputation "
        "pipeline. For Q1 journal positioning, the strongest defensible claim is that the final "
        "validation-calibrated MAT ensemble is the overall winner across the tested models and "
        "horizons, with statistically significant superiority under combined seed-level testing."
    )

    doc.add_heading("References", level=1)
    references = [
        "Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, L., and Polosukhin, I. (2017). Attention is all you need. Advances in Neural Information Processing Systems, 30.",
        "Che, Z., Purushotham, S., Cho, K., Sontag, D., and Liu, Y. (2018). Recurrent neural networks for multivariate time series with missing values. Scientific Reports, 8, 6085.",
        "Du, W., Cote, D., and Liu, Y. (2023). SAITS: Self-attention-based imputation for time series. Expert Systems with Applications, 219, 119619.",
        "Zeng, A., Chen, M., Zhang, L., and Xu, Q. (2023). Are Transformers effective for time series forecasting? Proceedings of the AAAI Conference on Artificial Intelligence, 37(9), 11121-11128.",
        "Nie, Y., Nguyen, N. H., Sinthong, P., and Kalagnanam, J. (2023). A time series is worth 64 words: Long-term forecasting with Transformers. International Conference on Learning Representations.",
        "van Buuren, S., and Groothuis-Oudshoorn, K. (2011). mice: Multivariate imputation by chained equations in R. Journal of Statistical Software, 45(3), 1-67.",
        "Cao, W., Wang, D., Li, J., Zhou, H., Li, L., and Li, Y. (2018). BRITS: Bidirectional recurrent imputation for time series. Advances in Neural Information Processing Systems, 31.",
        "Diebold, F. X., and Mariano, R. S. (1995). Comparing predictive accuracy. Journal of Business and Economic Statistics, 13(3), 253-263.",
    ]
    for i, ref in enumerate(references, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"[{i}] {ref}")
        set_run_font(r, size=10)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
